#!/usr/bin/env python3
import io
import json
import math
import re
import sys
import tempfile
from datetime import datetime, timezone
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


def clean(value, fallback="-"):
    if value is None:
        return fallback
    text = str(value).strip()
    return text if text else fallback


def parse_dt(value):
    if not value:
        return None

    text = str(value).strip()

    try:
        if text.endswith("Z"):
            text = text[:-1] + "+00:00"
        return datetime.fromisoformat(text).astimezone(timezone.utc)
    except Exception:
        return None


def format_dt_br(value):
    dt = parse_dt(value)
    if not dt:
        return clean(value)
    return dt.strftime("%d/%m/%Y %H:%M:%S")


def br_number(value, decimals=2):
    try:
        number = float(value)
        text = f"{number:.{decimals}f}".replace(".", ",")
    except Exception:
        return "-"

    if decimals <= 0:
        return text

    return text.rstrip("0").rstrip(",")


def format_value(value, unit):
    if value is None:
        return "-"

    try:
        number = float(value)
    except Exception:
        return clean(value)

    unit = clean(unit, "").lower()

    if unit == "bps":
        if abs(number) >= 1_000_000_000:
            return f"{br_number(number / 1_000_000_000)} Gbit/s"
        if abs(number) >= 1_000_000:
            return f"{br_number(number / 1_000_000)} Mbit/s"
        if abs(number) >= 1_000:
            return f"{br_number(number / 1_000)} Kbit/s"
        return f"{br_number(number, 0)} bit/s"

    if unit == "ms":
        return f"{br_number(number)} ms"

    if unit == "%":
        return f"{br_number(number, 3)} %"

    if unit == "d":
        days = int(number)
        hours = int((number - days) * 24)
        return f"{days}d {hours}h"

    return br_number(number)


def xml_text(element):
    """
    Extrai texto de um nó XML do python-docx sem usar xpath(namespaces=...),
    porque BaseOxmlElement.xpath() não aceita esse argumento em algumas versões.
    """
    values = []

    for node in element.iter():
        tag = getattr(node, "tag", "")

        if tag == qn("w:t") or tag.endswith("}t") or tag.endswith("}instrText"):
            if node.text:
                values.append(node.text)

    return " ".join(values)


def dedupe_cover_text_before_break(doc):
    """
    Remove parágrafos de texto duplicados na capa oficial antes da primeira
    quebra de página, preservando imagens/desenhos da capa.
    """
    body = doc._body._body
    children = list(body)

    break_index = None

    for index, child in enumerate(children):
        xml = child.xml
        if 'w:type="page"' in xml or "w:type='page'" in xml:
            break_index = index
            break

    if break_index is None:
        return

    seen = set()

    for child in children[:break_index + 1]:
        value = " ".join(xml_text(child).split())

        if not value:
            continue

        # Só remove duplicatas textuais. Não toca em imagens/desenhos.
        if value in seen:
            parent = child.getparent()
            if parent is not None:
                parent.remove(child)
        else:
            seen.add(value)


def purge_official_body_after_cover(doc):
    """
    Mantém a capa oficial real e remove TODO o miolo oficial antigo.

    O DOCX oficial gerado tem a capa na primeira página e depois uma quebra
    de página para o relatório de consumo. Antes, removíamos a partir de
    marcadores como "Contrato" ou "Consumo do Link"; isso deixava sobras como
    "UN – UNITINS..." no começo do documento.

    Agora a regra é mais segura:
    - encontra a primeira quebra de página;
    - preserva tudo até essa quebra, inclusive;
    - remove tudo depois dela, exceto sectPr;
    - o anexo técnico entra logo depois da capa.
    """
    body = doc._body._body
    children = list(body)

    break_index = None

    for index, child in enumerate(children):
        xml = child.xml
        if 'w:type="page"' in xml or "w:type='page'" in xml:
            break_index = index
            break

    if break_index is None:
        # Se não encontrar quebra, cai para o comportamento antigo por segurança.
        markers = [
            "UN –",
            "UN -",
            "Consumo do Link",
            "Gráfico de tráfego",
            "Contrato:",
            "Endereço:",
            "Banda Contratada",
        ]

        for index, child in enumerate(children):
            text = xml_text(child)
            if any(marker in text for marker in markers):
                break_index = index - 1
                break

    if break_index is None:
        doc.add_page_break()
        return

    # Remove tudo depois da quebra da capa.
    for child in children[break_index + 1:]:
        if child.tag.endswith("sectPr"):
            continue

        parent = child.getparent()
        if parent is not None:
            parent.remove(child)


def first_or_add_paragraph(container):
    if getattr(container, "paragraphs", None) and len(container.paragraphs) > 0:
        return container.paragraphs[0]
    return container.add_paragraph()


def tag_endswith(node, suffix):
    tag = getattr(node, "tag", "")
    return tag.endswith("}" + suffix) or tag == suffix


def closest_ancestor(node, suffix):
    current = node

    while current is not None:
        if tag_endswith(current, suffix):
            return current
        current = current.getparent()

    return None


def fix_footer_site_clipping(doc):
    """
    O texto www.novatelecom.com.br existe completo no rodapé,
    mas a caixa oficial fica perto demais da borda direita e o Word/LibreOffice
    corta o final ".br". Ajustamos só essa caixa.
    """
    footers = []

    for section in doc.sections:
        for footer_attr in ("footer", "first_page_footer", "even_page_footer"):
            footer = getattr(section, footer_attr, None)
            if footer is not None:
                footers.append(footer)

    for footer in footers:
        root = footer._element

        for node in list(root.iter()):
            if not tag_endswith(node, "t"):
                continue

            value = node.text or ""

            if "novatelecom.com" not in value:
                continue

            node.text = "www.novatelecom.com.br"

            paragraph = closest_ancestor(node, "p")
            if paragraph is not None:
                for sz in paragraph.iter():
                    if tag_endswith(sz, "sz"):
                        sz.set(qn("w:val"), "20")

            anchor = closest_ancestor(node, "anchor")
            if anchor is not None:
                for position_h in anchor.iter():
                    if not tag_endswith(position_h, "positionH"):
                        continue

                    for pos_offset in position_h.iter():
                        if tag_endswith(pos_offset, "posOffset"):
                            # Move a caixa aproximadamente 0,4 cm para a esquerda.
                            pos_offset.text = "4200000"
                            break


def set_section_layout(doc):
    """
    Ajusta apenas a área útil da seção final, sem clonar/recortar cabeçalho.
    O cabeçalho/rodapé real permanece vindo do documento oficial base.
    """
    section = doc.sections[-1]
    section.top_margin = Inches(1.96)
    section.bottom_margin = Inches(0.95)
    section.left_margin = Inches(0.72)
    section.right_margin = Inches(0.82)
    return section


def block_kind(block):
    text = " ".join([
        clean(block.get("title"), ""),
        clean(block.get("description"), ""),
        clean(block.get("sensorType"), ""),
        clean(block.get("unit"), ""),
        " ".join(clean(s.get("label") or s.get("name"), "") for s in block.get("series", [])),
    ]).lower()

    if "ping" in text or "latência" in text or "latencia" in text:
        return "ping"
    if "uptime" in text or "atividade" in text or "disponibilidade" in text:
        return "uptime"
    if "traffic" in text or "tráfego" in text or "trafego" in text or "link" in text or "interface" in text:
        return "traffic"
    return "other"


def chart_label(block):
    kind = block_kind(block)
    if kind == "ping":
        return "Ping"
    if kind == "uptime":
        return "System Uptime"
    if kind == "traffic":
        return "Link Traffic"
    return clean(block.get("title"), "Sensor")


def order_blocks(blocks):
    order = {"ping": 0, "uptime": 1, "traffic": 2}
    return sorted(blocks, key=lambda item: order.get(block_kind(item), 9))


def series_stats(series):
    return series.get("stats") or {}


def numeric_points(series):
    points = []

    for point in series.get("points", []) or []:
        dt = parse_dt(point.get("timestamp"))
        value = point.get("value")

        if not dt or value is None:
            continue

        try:
            number = float(value)
        except Exception:
            continue

        if math.isfinite(number):
            points.append((dt, number))

    return points


def packet_loss_series(block):
    for series in block.get("series", []) or []:
        label = f"{series.get('label')} {series.get('name')} {series.get('kind')}".lower()
        unit = clean(series.get("unit"), "").lower()

        if unit == "%" or "perda" in label or "loss" in label:
            return series

    return None


def first_series_by_unit(block, unit):
    for series in block.get("series", []) or []:
        if clean(series.get("unit"), "").lower() == unit:
            return series
    series = block.get("series") or []
    return series[0] if series else {}


def data_period(block):
    timestamps = []
    for series in block.get("series", []) or []:
        for point in series.get("points", []) or []:
            timestamp = point.get("timestamp")
            if timestamp:
                timestamps.append(timestamp)

    if len(timestamps) < 2:
        return "-"

    timestamps.sort()
    return f"{format_dt_br(timestamps[0])} - {format_dt_br(timestamps[-1])}"


def data_duration_seconds(block):
    dates = []

    for series in block.get("series", []) or []:
        for point in series.get("points", []) or []:
            dt = parse_dt(point.get("timestamp"))
            if dt:
                dates.append(dt.timestamp())

    dates.sort()

    if len(dates) < 2:
        return 0

    return max(0, dates[-1] - dates[0])


def format_percent(value):
    return f"{br_number(value, 3)} %"


def format_duration(total_seconds):
    seconds = max(0, int(round(total_seconds)))
    days = seconds // 86400
    hours = (seconds % 86400) // 3600
    minutes = (seconds % 3600) // 60
    sec = seconds % 60
    return f"{days:02d}d {hours:02d}h {minutes:02d}m {sec:02d}s"


def request_counts(block):
    loss = packet_loss_series(block)

    if loss:
        points = [point for point in loss.get("points", []) or [] if point.get("value") is not None]
        if points:
            failed = sum(1 for point in points if float(point.get("value") or 0) > 0)
            good = max(0, len(points) - failed)
            return good, failed, len(points)

    counts = [
        int((series.get("stats") or {}).get("points") or len(series.get("points", []) or []) or 0)
        for series in block.get("series", []) or []
    ]
    total = max(counts) if counts else 0
    return total, 0, total


def availability_label(block):
    good, failed, total = request_counts(block)
    if not total:
        return "-"

    duration = data_duration_seconds(block)
    good_ratio = good / total
    failed_ratio = failed / total

    return (
        f"OK: {format_percent(good_ratio * 100)} [{format_duration(duration * good_ratio)}] "
        f"Inoperante: {format_percent(failed_ratio * 100)} [{format_duration(duration * failed_ratio)}]"
    )


def request_label(block):
    good, failed, total = request_counts(block)
    if not total:
        return "-"

    return f"Bom: {format_percent(good / total * 100)} [{good}] Falha: {format_percent(failed / total * 100)} [{failed}]"


def period_row(payload, report):
    period = report.get("period") or payload.get("period") or {}
    start = period.get("from") or payload.get("periodFrom") or payload.get("from")
    finish = period.get("to") or payload.get("periodTo") or payload.get("to")

    if start and finish:
        return f"{format_dt_br(start)} - {format_dt_br(finish)}"

    label = clean(period.get("label") or payload.get("periodLabel") or payload.get("collectionPeriodLabel"), "")
    return label if label else "-"



def avg_numeric(series):
    if not series:
        return None

    stats = series_stats(series)
    value = stats.get("avg")

    if value is not None:
        try:
            return float(value)
        except Exception:
            pass

    pts = numeric_points(series)

    if not pts:
        return None

    return sum(value for _, value in pts) / len(pts)


def max_numeric(series):
    if not series:
        return None

    stats = series_stats(series)
    value = stats.get("max")

    if value is not None:
        try:
            return float(value)
        except Exception:
            pass

    pts = numeric_points(series)

    if not pts:
        return None

    return max(value for _, value in pts)


def traffic_series(block):
    download = None
    upload = None
    total = None

    for series in block.get("series", []) or []:
        label = f"{series.get('label')} {series.get('name')} {series.get('kind')}".lower()
        unit = clean(series.get("unit"), "").lower()

        if unit != "bps":
            continue

        if any(term in label for term in ["receb", "download", "entrada", "inbound", "receive"]):
            download = series
        elif any(term in label for term in ["envi", "upload", "saida", "saída", "outbound", "send"]):
            upload = series
        elif any(term in label for term in ["total", "traffic", "tráfego", "trafego"]):
            total = series

    bps_series = [
        series
        for series in block.get("series", []) or []
        if clean(series.get("unit"), "").lower() == "bps"
    ]

    if download is None and bps_series:
        download = bps_series[0]

    if upload is None and len(bps_series) > 1:
        upload = bps_series[1]

    if total is None and len(bps_series) > 2:
        total = bps_series[2]

    return download, upload, total


def parse_bandwidth_to_bps(value):
    if value is None:
        return None

    text_value = str(value).strip().lower().replace(",", ".")

    match = re.search(r"([0-9]+(?:\.[0-9]+)?)\s*(gbit/s|gbps|g\b|mbit/s|mbps|m\b|kbit/s|kbps|k\b)", text_value)

    if not match:
        return None

    number = float(match.group(1))
    unit = match.group(2)

    if unit.startswith("g"):
        return number * 1_000_000_000

    if unit.startswith("m"):
        return number * 1_000_000

    if unit.startswith("k"):
        return number * 1_000

    return None


def infer_capacity_bps(block):
    candidates = []

    for source in [
        block,
        block.get("consumption") or {},
        block.get("contract") or {},
        block.get("link") or {},
    ]:
        if not isinstance(source, dict):
            continue

        for key, value in source.items():
            key_l = str(key).lower()

            if not any(term in key_l for term in ["capacity", "capacidade", "bandwidth", "banda", "contract"]):
                continue

            parsed = parse_bandwidth_to_bps(value)

            if parsed:
                candidates.append(parsed)

    blob = " ".join(str(value) for value in [
        block.get("capacity"),
        block.get("capacityLabel"),
        block.get("bandwidth"),
        block.get("bandwidthLabel"),
        block.get("contractedBandwidth"),
        block.get("contractedBandwidthLabel"),
        block.get("description"),
        block.get("sensorType"),
        (block.get("consumption") or {}).get("capacityLabel"),
        (block.get("consumption") or {}).get("bandwidthLabel"),
    ] if value)

    parsed = parse_bandwidth_to_bps(blob)

    if parsed:
        candidates.append(parsed)

    if candidates:
        return max(candidates)

    download, upload, total = traffic_series(block)
    peaks = []

    for series in [download, upload, total]:
        peak = max_numeric(series)

        if peak:
            peaks.append(peak)

    if not peaks:
        return None

    observed_peak = max(peaks)

    if observed_peak <= 0:
        return None

    mbps = observed_peak / 1_000_000

    for capacity in [10, 20, 30, 40, 50, 100, 200, 300, 500, 1000]:
        if mbps <= capacity * 0.92:
            return capacity * 1_000_000

    return None


def utilization_label(value_bps, capacity_bps):
    if value_bps is None or not capacity_bps:
        return "-"

    return f"{br_number((value_bps / capacity_bps) * 100, 2)}%"


def summary_rows(payload, report, block):
    kind = block_kind(block)
    primary = first_series_by_unit(block, "ms" if kind == "ping" else "d" if kind == "uptime" else "bps")
    stats = series_stats(primary)

    if kind == "traffic":
        download, upload, total = traffic_series(block)
        capacity_bps = infer_capacity_bps(block)

        avg_down = avg_numeric(download) if download else None
        avg_up = avg_numeric(upload) if upload else None
        peak_total = max([
            value
            for value in [
                max_numeric(download) if download else None,
                max_numeric(upload) if upload else None,
                max_numeric(total) if total else None,
            ]
            if value is not None
        ], default=None)

        avg_total = None

        if avg_down is not None or avg_up is not None:
            avg_total = (avg_down or 0) + (avg_up or 0)
        elif total:
            avg_total = avg_numeric(total)

        return [
            ["Recebido (média)", f"{format_value(avg_down, 'bps')} ({utilization_label(avg_down, capacity_bps)})"],
            ["Enviado (média)", f"{format_value(avg_up, 'bps')} ({utilization_label(avg_up, capacity_bps)})"],
            ["Utilização média", utilization_label(avg_total, capacity_bps)],
            ["Pico total", f"{format_value(peak_total, 'bps')} ({utilization_label(peak_total, capacity_bps)})"],
            ["Capacidade", format_value(capacity_bps, "bps") if capacity_bps else "-"],
            ["Período do relatório", period_row(payload, report)],
            ["Tipo de sensor", clean(block.get("sensorType"), chart_label(block))],
            ["Sonda, grupo, dispositivo", clean(block.get("probePath") or report.get("hostName") or (report.get("host") or {}).get("hostName") or (report.get("unit") or {}).get("name"))],
        ]

    rows = [
        ["Período do relatório", period_row(payload, report)],
        ["Período com dados", data_period(block)],
        ["Horas de relatório", "24 / 7"],
        ["Tipo de sensor", clean(block.get("sensorType"), chart_label(block))],
        ["Sonda, grupo, dispositivo", clean(block.get("probePath") or report.get("hostName") or (report.get("host") or {}).get("hostName") or (report.get("unit") or {}).get("name"))],
        ["Estatísticas de tempo de atividade", availability_label(block)],
        ["Estatísticas de solicitação", request_label(block)],
    ]

    if kind == "ping":
        loss = packet_loss_series(block)
        rows.extend([
            ["Média (Tempo de ping)", format_value(stats.get("avg"), primary.get("unit", "ms"))],
            ["Mínimo", format_value(stats.get("min"), primary.get("unit", "ms"))],
            ["Máximo", format_value(stats.get("max"), primary.get("unit", "ms"))],
        ])
        if loss:
            rows.append(["Perda de pacote média", format_value(series_stats(loss).get("avg"), loss.get("unit", "%"))])
    elif kind == "uptime":
        rows.extend([
            ["Média (Tempo de atividade do sistema)", format_value(stats.get("avg"), primary.get("unit", "d"))],
            ["Mínimo", format_value(stats.get("min"), primary.get("unit", "d"))],
            ["Máximo", format_value(stats.get("max"), primary.get("unit", "d"))],
        ])

    return rows


def channel_rows(block):
    kind = block_kind(block)

    if kind == "traffic":
        download, upload, total = traffic_series(block)
        capacity_bps = infer_capacity_bps(block)

        rows = [["Canal", "Último", "Mín.", "Máx.", "Média", "Util."]]

        for label, series in [
            ("Tráfego recebido", download),
            ("Tráfego enviado", upload),
        ]:
            if not series:
                continue

            stats = series_stats(series)
            avg_value = stats.get("avg")
            rows.append([
                label,
                format_value(stats.get("last") or stats.get("current") or avg_value, series.get("unit")),
                format_value(stats.get("min"), series.get("unit")),
                format_value(stats.get("max"), series.get("unit")),
                format_value(avg_value, series.get("unit")),
                utilization_label(avg_value, capacity_bps),
            ])

        if len(rows) == 1:
            consumption = block.get("consumption") or {}
            rows.extend([
                ["Tráfego recebido", clean(consumption.get("avgReceiveLabel")), "-", clean(consumption.get("peakReceiveLabel")), clean(consumption.get("avgReceiveLabel")), "-"],
                ["Tráfego enviado", clean(consumption.get("avgSendLabel")), "-", clean(consumption.get("peakSendLabel")), clean(consumption.get("avgSendLabel")), "-"],
            ])

        return rows

    rows = [["Canal", "Média", "Mín.", "Máx."]]
    series_list = block.get("series") or []

    if kind == "ping":
        series_list = sorted(series_list, key=lambda series: 0 if clean(series.get("unit"), "").lower() == "ms" else 1)

    for series in series_list[:6]:
        stats = series_stats(series)
        rows.append([
            clean(series.get("label") or series.get("name"), "Canal"),
            format_value(stats.get("avg"), series.get("unit")),
            format_value(stats.get("min"), series.get("unit")),
            format_value(stats.get("max"), series.get("unit")),
        ])

    return rows


def bucket_traffic_points(download_pts, upload_pts, max_buckets=26):
    """
    Reduz muitos pontos de tráfego para buckets visuais.
    Usa pico por bucket para manter eventos relevantes e deixar barras legíveis.
    """
    if not download_pts or len(download_pts) <= max_buckets:
        return download_pts, upload_pts

    upload_map = {dt: value for dt, value in upload_pts or []}
    step = max(1, math.ceil(len(download_pts) / max_buckets))

    down_bucketed = []
    up_bucketed = []

    for index in range(0, len(download_pts), step):
        group = download_pts[index:index + step]

        if not group:
            continue

        dt = group[len(group) // 2][0]
        down_peak = max(value for _, value in group)
        up_values = [upload_map.get(point_dt, 0) for point_dt, _ in group]
        up_peak = max(up_values) if up_values else 0

        down_bucketed.append((dt, down_peak))
        up_bucketed.append((dt, up_peak))

    return down_bucketed, up_bucketed



def render_chart_png(block):
    import matplotlib
    matplotlib.use("Agg")
    import matplotlib.dates as mdates
    import matplotlib.pyplot as plt

    kind = block_kind(block)

    def style_axis(axis, show_x=True):
        axis.set_facecolor("#ffffff")
        axis.grid(True, color="#e8eef5", linewidth=0.48)
        axis.tick_params(axis="both", labelsize=6.6, colors="#4b5563")
        axis.spines["top"].set_visible(False)
        axis.spines["right"].set_visible(False)
        axis.spines["left"].set_color("#c7d0da")
        axis.spines["bottom"].set_color("#c7d0da")
        if not show_x:
            axis.tick_params(axis="x", labelbottom=False)

    def apply_date_axis(axis):
        axis.xaxis.set_major_locator(mdates.AutoDateLocator(minticks=4, maxticks=6))
        axis.xaxis.set_major_formatter(mdates.DateFormatter("%d/%m %H:%M"))

    def save_figure(fig):
        buffer = io.BytesIO()
        fig.savefig(buffer, format="png", bbox_inches="tight", pad_inches=0.10)
        plt.close(fig)
        return buffer.getvalue()

    if kind == "ping":
        ping = first_series_by_unit(block, "ms")
        loss = packet_loss_series(block)

        ping_pts = numeric_points(ping)
        loss_pts = numeric_points(loss) if loss else []

        has_loss = bool(loss_pts)

        if has_loss:
            fig, axes = plt.subplots(
                2,
                1,
                figsize=(9.65, 4.35),
                dpi=230,
                sharex=True,
                gridspec_kw={"height_ratios": [3.25, 1.05], "hspace": 0.08},
            )
            ax = axes[0]
            loss_ax = axes[1]
        else:
            fig, ax = plt.subplots(figsize=(9.65, 3.70), dpi=230)
            loss_ax = None

        fig.patch.set_facecolor("white")
        style_axis(ax, show_x=not has_loss)

        plotted = False

        if ping_pts:
            x = [item[0] for item in ping_pts]
            y = [item[1] for item in ping_pts]

            ax.plot(
                x,
                y,
                color="#1d4ed8",
                linewidth=1.22,
                label="Tempo de ping",
                solid_capstyle="round",
            )
            ax.fill_between(x, y, color="#1d4ed8", alpha=0.045)
            ax.set_ylim(bottom=0, top=max(5, max(y) * 1.18 if y else 5))
            ax.legend(loc="upper left", fontsize=6.8, frameon=False)
            plotted = True

        ax.set_ylabel("ms", fontsize=7.8, color="#374151")
        ax.set_title(chart_label(block), fontsize=8.6, color="#111827", pad=6)

        if loss_ax is not None:
            style_axis(loss_ax, show_x=True)

            x_loss = [item[0] for item in loss_pts]
            y_loss = [item[1] for item in loss_pts]

            loss_ax.plot(
                x_loss,
                y_loss,
                color="#f97316",
                linewidth=0.72,
                alpha=0.55,
                label="Perda de pacote",
                drawstyle="steps-mid",
            )

            spikes = [(x, y) for x, y in loss_pts if y and y > 0]
            if spikes:
                loss_ax.vlines(
                    [x for x, _ in spikes],
                    0,
                    [y for _, y in spikes],
                    color="#f97316",
                    linewidth=0.85,
                    alpha=0.58,
                )

            loss_ax.set_ylim(-2, 102)
            loss_ax.set_ylabel("%", fontsize=7.8, color="#374151")
            loss_ax.legend(loc="upper left", fontsize=6.8, frameon=False)
            apply_date_axis(loss_ax)
            fig.autofmt_xdate(rotation=22, ha="right")
            plotted = True
        else:
            apply_date_axis(ax)
            fig.autofmt_xdate(rotation=22, ha="right")

        if not plotted:
            ax.text(
                0.5,
                0.5,
                "Sem dados para o gráfico",
                transform=ax.transAxes,
                ha="center",
                va="center",
                fontsize=9,
                color="#6b7280",
            )

        return save_figure(fig)

    if kind == "traffic":
        fig, ax = plt.subplots(figsize=(9.65, 4.05), dpi=230)
        fig.patch.set_facecolor("white")
        style_axis(ax)

        download, upload, total = traffic_series(block)
        capacity_bps = infer_capacity_bps(block)

        download_pts_raw = numeric_points(download) if download else []
        upload_pts_raw = numeric_points(upload) if upload else []
        download_pts, upload_pts = bucket_traffic_points(download_pts_raw, upload_pts_raw)

        plotted = []

        values_for_ylim = []

        if download_pts:
            x_down = [item[0] for item in download_pts]
            y_down = [item[1] / 1_000_000 for item in download_pts]
            values_for_ylim.extend(y_down)

            ax.plot(
                x_down,
                y_down,
                color="#1d4ed8",
                linewidth=1.28,
                label="Tráfego recebido",
                solid_capstyle="round",
            )
            ax.fill_between(x_down, y_down, color="#1d4ed8", alpha=0.055)
            plotted.append("Tráfego recebido")

        if upload_pts:
            x_up = [item[0] for item in upload_pts]
            y_up = [item[1] / 1_000_000 for item in upload_pts]
            values_for_ylim.extend(y_up)

            ax.plot(
                x_up,
                y_up,
                color="#65a30d",
                linewidth=1.16,
                linestyle=(0, (4, 2)),
                label="Tráfego enviado",
                solid_capstyle="round",
            )
            ax.fill_between(x_up, y_up, color="#65a30d", alpha=0.040)
            plotted.append("Tráfego enviado")

        if not plotted and total:
            pts = numeric_points(total)
            if pts:
                x_total = [item[0] for item in pts]
                y_total = [item[1] / 1_000_000 for item in pts]
                values_for_ylim.extend(y_total)

                ax.plot(
                    x_total,
                    y_total,
                    color="#1d4ed8",
                    linewidth=1.18,
                    label="Tráfego total",
                    solid_capstyle="round",
                )
                plotted.append("Tráfego total")

        if capacity_bps:
            capacity_mbps = capacity_bps / 1_000_000
            values_for_ylim.append(capacity_mbps)
            ax.axhline(
                capacity_mbps,
                color="#ef4444",
                linestyle=(0, (5, 4)),
                linewidth=0.82,
                alpha=0.58,
                label=f"Capacidade ({br_number(capacity_mbps, 0)} Mbps)",
            )
            plotted.append("Capacidade")

        if values_for_ylim:
            top = max(values_for_ylim)
            ax.set_ylim(bottom=0, top=max(1, top * 1.12))

        ax.set_ylabel("Mbps", fontsize=7.8, color="#374151")
        ax.set_title(chart_label(block), fontsize=8.6, color="#111827", pad=6)

        apply_date_axis(ax)
        fig.autofmt_xdate(rotation=22, ha="right")

        if plotted:
            ax.legend(
                loc="upper center",
                bbox_to_anchor=(0.5, -0.18),
                ncol=min(3, len(plotted)),
                fontsize=6.8,
                frameon=False,
                handlelength=2.8,
                columnspacing=1.4,
            )
        else:
            ax.text(
                0.5,
                0.5,
                "Sem dados para o gráfico",
                transform=ax.transAxes,
                ha="center",
                va="center",
                fontsize=9,
                color="#6b7280",
            )

        return save_figure(fig)

    fig, ax = plt.subplots(figsize=(9.65, 4.05), dpi=230)
    fig.patch.set_facecolor("white")
    style_axis(ax)

    plotted = []

    if kind == "uptime":
        uptime = first_series_by_unit(block, "d")
        pts = numeric_points(uptime)

        if pts:
            ax.plot(
                [x for x, _ in pts],
                [y for _, y in pts],
                color="#1d4ed8",
                linewidth=1.16,
                label="Tempo de atividade do sistema",
            )
            ax.fill_between([x for x, _ in pts], [y for _, y in pts], alpha=0.08, color="#1d4ed8")
            plotted.append("Tempo de atividade do sistema")

        ax.set_ylabel("dias", fontsize=7.8, color="#374151")

    else:
        colors = ["#1d4ed8", "#65a30d", "#7c3aed", "#f97316"]
        idx = 0

        for series in block.get("series", []) or []:
            pts = numeric_points(series)
            if not pts:
                continue

            unit = clean(series.get("unit"), "").lower()
            values = [y for _, y in pts]

            if unit == "bps":
                values = [value / 1_000_000 for value in values]
                ax.set_ylabel("Mbps", fontsize=7.8, color="#374151")
            else:
                ax.set_ylabel(unit or "valor", fontsize=7.8, color="#374151")

            label = clean(series.get("label") or series.get("name"), "Série")
            ax.plot([x for x, _ in pts], values, color=colors[idx % len(colors)], linewidth=1.10, label=label)
            plotted.append(label)
            idx += 1

    if plotted:
        ax.legend(
            loc="upper center",
            bbox_to_anchor=(0.5, -0.18),
            ncol=min(3, len(plotted)),
            fontsize=6.8,
            frameon=False,
        )

    ax.set_title(chart_label(block), fontsize=8.6, color="#111827", pad=6)
    apply_date_axis(ax)
    fig.autofmt_xdate(rotation=22, ha="right")

    if not plotted:
        ax.text(
            0.5,
            0.5,
            "Sem dados para o gráfico",
            transform=ax.transAxes,
            ha="center",
            va="center",
            fontsize=9,
            color="#6b7280",
        )

    return save_figure(fig)

def clear_cell(cell):
    cell.text = ""
    for paragraph in list(cell.paragraphs):
        p = paragraph._element
        p.getparent().remove(p)


def set_cell_border_color(cell, color="D9E2EC"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")

    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)

    for edge in ["top", "left", "bottom", "right"]:
        tag = "w:" + edge
        element = borders.find(qn(tag))

        if element is None:
            element = OxmlElement(tag)
            borders.append(element)

        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:color"), color)


def traffic_kpi_cards(block):
    download, upload, total = traffic_series(block)
    capacity_bps = infer_capacity_bps(block)

    avg_down = avg_numeric(download) if download else None
    avg_up = avg_numeric(upload) if upload else None

    avg_total = None

    if avg_down is not None or avg_up is not None:
        avg_total = (avg_down or 0) + (avg_up or 0)
    elif total:
        avg_total = avg_numeric(total)

    peak_total = max([
        value
        for value in [
            max_numeric(download) if download else None,
            max_numeric(upload) if upload else None,
            max_numeric(total) if total else None,
        ]
        if value is not None
    ], default=None)

    return [
        ("Recebido", format_value(avg_down, "bps"), utilization_label(avg_down, capacity_bps)),
        ("Enviado", format_value(avg_up, "bps"), utilization_label(avg_up, capacity_bps)),
        ("Utilização", utilization_label(avg_total, capacity_bps), "média"),
        ("Pico total", format_value(peak_total, "bps"), utilization_label(peak_total, capacity_bps)),
    ]




def sensor_kpi_cards(block):
    kind = block_kind(block)

    if kind == "traffic":
        download, upload, total = traffic_series(block)
        capacity_bps = infer_capacity_bps(block)

        avg_down = avg_numeric(download)
        avg_up = avg_numeric(upload)

        avg_total = None
        if avg_down is not None or avg_up is not None:
            avg_total = (avg_down or 0) + (avg_up or 0)
        elif total:
            avg_total = avg_numeric(total)

        peak_total = max([
            value for value in [
                max_numeric(download),
                max_numeric(upload),
                max_numeric(total),
            ] if value is not None
        ], default=None)

        return [
            ("Recebido", format_value(avg_down, "bps"), utilization_label(avg_down, capacity_bps)),
            ("Enviado", format_value(avg_up, "bps"), utilization_label(avg_up, capacity_bps)),
            ("Utilização", utilization_label(avg_total, capacity_bps), "média"),
            ("Pico total", format_value(peak_total, "bps"), utilization_label(peak_total, capacity_bps)),
        ]

    if kind == "ping":
        ping = first_series_by_unit(block, "ms")
        loss = packet_loss_series(block)
        ping_stats = series_stats(ping)
        loss_stats = series_stats(loss) if loss else {}

        return [
            ("Média", format_value(ping_stats.get("avg"), "ms"), "tempo de ping"),
            ("Mínimo", format_value(ping_stats.get("min"), "ms"), "melhor valor"),
            ("Máximo", format_value(ping_stats.get("max"), "ms"), "pior valor"),
            ("Perda", format_value(loss_stats.get("avg"), "%") if loss else "-", "pacote"),
        ]

    if kind == "uptime":
        uptime = first_series_by_unit(block, "d")
        stats = series_stats(uptime)

        return [
            ("Média", format_value(stats.get("avg"), "d"), "uptime"),
            ("Mínimo", format_value(stats.get("min"), "d"), "registrado"),
            ("Máximo", format_value(stats.get("max"), "d"), "registrado"),
            ("Status", "OK", "operacional"),
        ]

    primary = (block.get("series") or [{}])[0]
    stats = series_stats(primary)
    return [
        ("Média", format_value(stats.get("avg"), primary.get("unit")), "sensor"),
        ("Mínimo", format_value(stats.get("min"), primary.get("unit")), "sensor"),
        ("Máximo", format_value(stats.get("max"), primary.get("unit")), "sensor"),
        ("Status", "OK", "geral"),
    ]


def clear_cell(cell):
    cell.text = ""
    for paragraph in list(cell.paragraphs):
        p = paragraph._element
        p.getparent().remove(p)


def set_cell_margins(cell, top=70, start=80, bottom=70, end=80):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")

    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)

    for margin_name, value in [
        ("top", top),
        ("start", start),
        ("bottom", bottom),
        ("end", end),
    ]:
        node = tc_mar.find(qn("w:" + margin_name))
        if node is None:
            node = OxmlElement("w:" + margin_name)
            tc_mar.append(node)

        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_border_color(cell, color="D8E3EF", size="4"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")

    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)

    for edge in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        tag = "w:" + edge
        element = borders.find(qn(tag))

        if element is None:
            element = OxmlElement(tag)
            borders.append(element)

        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:color"), color)


def add_kpi_cards(doc, cards):
    table = doc.add_table(rows=1, cols=len(cards))
    table.autofit = True

    accents = ["EAF2FF", "ECFDF3", "F4ECFF", "FFF3E8"]
    border_colors = ["B7C9EA", "BDE5C8", "D8C2F0", "F3C99E"]
    value_colors = [
        RGBColor(18, 58, 90),
        RGBColor(29, 100, 45),
        RGBColor(91, 38, 140),
        RGBColor(201, 78, 15),
    ]

    for index, (label, value, detail) in enumerate(cards):
        cell = table.cell(0, index)
        clear_cell(cell)
        set_cell_shading(cell, accents[index % len(accents)])
        set_cell_border_color(cell, border_colors[index % len(border_colors)], "6")
        set_cell_margins(cell, top=125, start=115, bottom=125, end=115)

        p_label = cell.add_paragraph()
        p_label.paragraph_format.space_before = Pt(0)
        p_label.paragraph_format.space_after = Pt(0)
        p_label.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_label = p_label.add_run(label)
        r_label.font.name = "Arial"
        r_label.font.size = Pt(8.0)
        r_label.font.color.rgb = RGBColor(70, 80, 92)
        r_label.bold = True

        p_value = cell.add_paragraph()
        p_value.paragraph_format.space_before = Pt(0)
        p_value.paragraph_format.space_after = Pt(0)
        p_value.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_value = p_value.add_run(clean(value))
        r_value.font.name = "Arial"
        r_value.font.size = Pt(18.4)
        r_value.font.color.rgb = value_colors[index % len(value_colors)]
        r_value.bold = True

        p_detail = cell.add_paragraph()
        p_detail.paragraph_format.space_before = Pt(0)
        p_detail.paragraph_format.space_after = Pt(0)
        p_detail.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_detail = p_detail.add_run(clean(detail))
        r_detail.font.name = "Arial"
        r_detail.font.size = Pt(7.6)
        r_detail.font.color.rgb = RGBColor(95, 105, 118)

    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_before = Pt(0)
    spacer.paragraph_format.space_after = Pt(0)

    return table

def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 0.98

    run = paragraph.add_run(clean(text, ""))
    run.font.name = "Arial"
    run.font.size = Pt(8.05 if not bold else 8.45)
    run.bold = bold
    run.font.color.rgb = RGBColor(31, 41, 55)

def add_table(doc, rows, col_count):
    table = doc.add_table(rows=len(rows), cols=col_count)
    table.autofit = True

    for row_index, row in enumerate(rows):
        for col_index in range(col_count):
            value = row[col_index] if col_index < len(row) else ""
            cell = table.cell(row_index, col_index)

            if row_index == 0:
                set_cell_shading(cell, "F2F6FA")
                set_cell_border_color(cell, "C9D5E3", "5")
            else:
                set_cell_shading(cell, "FFFFFF")
                set_cell_border_color(cell, "DDE6EF", "4")

            set_cell_margins(cell, top=75, start=90, bottom=75, end=90)
            set_cell_text(cell, value, bold=(row_index == 0 or col_index == 0))

    return table

def add_heading(doc, text, level=1):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(2 if level == 1 else 1)

    run = paragraph.add_run(text)
    run.bold = True
    run.font.name = "Arial"
    run.font.color.rgb = RGBColor(18, 58, 90)
    run.font.size = Pt(14.2 if level == 1 else 9.2)

    return paragraph

def add_small(doc, text):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(1)

    run = paragraph.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(7.5)
    run.font.color.rgb = RGBColor(80, 90, 100)

    return paragraph



def patch_header_footer_edge_bleed(docx_path, expand_emu=760000):
    """
    Ajuste final pós-save no DOCX.

    O fio branco da direita vem das artes ancoradas do header/footer técnico
    terminando alguns milímetros antes da borda. O python-docx nem sempre
    preserva/expõe todos os nós internos dessas artes agrupadas, então aqui
    abrimos o DOCX como zip e ampliamos diretamente os cx grandes nos XMLs
    de header/footer.

    760000 EMU ~= 0,83 cm. É intencionalmente maior que o fio visível para
    garantir sangria lateral, sem deslocar o conteúdo.
    """
    import os
    import re
    import shutil
    import tempfile
    import zipfile
    from pathlib import Path

    source = Path(docx_path)
    tmp_dir = Path(tempfile.mkdtemp(prefix="nova-docx-hf-"))
    work = tmp_dir / "work"
    work.mkdir(parents=True, exist_ok=True)

    try:
        with zipfile.ZipFile(source, "r") as zin:
            zin.extractall(work)

        word_dir = work / "word"
        xml_files = list(word_dir.glob("header*.xml")) + list(word_dir.glob("footer*.xml"))

        # Não alteramos header/footer muito pequenos; esses tendem a ser vazios
        # ou de capa. Só mexemos nos XMLs que carregam artes grandes.
        for xml_file in xml_files:
            xml = xml_file.read_text(encoding="utf-8", errors="ignore")

            if "wp:anchor" not in xml:
                continue

            def repl(match):
                name = match.group(1)
                value = int(match.group(2))

                # Expande apenas desenhos grandes. Textos/caixas pequenas ficam intactos.
                if value >= 4500000:
                    return f'{name}="{value + expand_emu}"'

                return match.group(0)

            patched = re.sub(r'(cx)="(\d+)"', repl, xml)

            if patched != xml:
                xml_file.write_text(patched, encoding="utf-8")

        tmp_docx = tmp_dir / source.name

        with zipfile.ZipFile(tmp_docx, "w", zipfile.ZIP_DEFLATED) as zout:
            for file in work.rglob("*"):
                if file.is_file():
                    zout.write(file, file.relative_to(work).as_posix())

        shutil.move(str(tmp_docx), str(source))
    finally:
        shutil.rmtree(tmp_dir, ignore_errors=True)



def build_docx(base_docx, payload, output_path):
    doc = Document(base_docx)

    fix_footer_site_clipping(doc)
    purge_official_body_after_cover(doc)
    set_section_layout(doc)

    month = clean(payload.get("monthSlashLabel") or payload.get("monthLabel"), "-")

    first_block = True

    with tempfile.TemporaryDirectory(prefix="nova-docx-assets-") as temp_dir:
        for report in payload.get("reports", []):
            unit = report.get("unit") or {}
            unit_name = clean(unit.get("name") or unit.get("label"), "Unidade")

            for block in order_blocks(report.get("blocks", [])):
                if not first_block:
                    doc.add_page_break()

                first_block = False

                add_heading(doc, f"{unit_name}: {chart_label(block)}", 1)

                add_heading(doc, "Resumo do sensor", 2)
                add_kpi_cards(doc, sensor_kpi_cards(block))

                add_heading(doc, "Gráfico do sensor", 2)
                try:
                    chart_png = render_chart_png(block)
                    safe_name = re.sub(r"[^a-zA-Z0-9_-]+", "_", chart_label(block)).strip("_") or "chart"
                    chart_path = Path(temp_dir) / f"{safe_name}.png"
                    chart_path.write_bytes(chart_png)
                    doc.add_picture(str(chart_path), width=Inches(7.15))
                except Exception as chart_error:
                    add_small(doc, f"Gráfico indisponível: {chart_error}")

                add_heading(doc, "Canal", 2)
                rows = channel_rows(block)
                add_table(doc, rows, len(rows[0]))

        doc.save(output_path)
    patch_header_footer_edge_bleed(output_path)


def main():
    if len(sys.argv) != 4:
        print("Uso: render_complete_from_official_docx.py base_official.docx payload.json output.docx", file=sys.stderr)
        return 2

    base_docx = Path(sys.argv[1])
    payload = json.loads(Path(sys.argv[2]).read_text())
    output_path = Path(sys.argv[3])

    build_docx(base_docx, payload, output_path)

    return 0


if __name__ == "__main__":
    raise SystemExit(main())
